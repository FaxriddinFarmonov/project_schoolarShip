import io
import deepl
import pytesseract
from PIL import Image
from docx import Document as DocxDocument
from docx.shared import Inches
from io import BytesIO
from django.http import HttpResponse
from django.shortcuts import get_object_or_404
from app.models.upload_file import UploadedFile

# 🔑 DeepL API kalitini kiriting
DEEPL_API_KEY = "77635c486f2deb7d880f8d757c342d45c83c02ac"
translator = deepl.Translator(DEEPL_API_KEY)

# 🔍 OCR uchun Tesseract sozlash (Windows uchun yo'l)
pytesseract.pytesseract.tesseract_cmd = r"C:\Program Files\Tesseract-OCR\tesseract.exe"

def translate_text_with_deepl(text, target_language="UZ"):
    """DeepL yordamida matn tarjima qiladi."""
    if not text.strip():
        return ""
    result = translator.translate_text(text, target_lang=target_language)
    return result.text

def extract_text_from_image(image_stream):
    """Rasmdagi matnni OCR orqali ajratib olish."""
    image = Image.open(image_stream)
    text = pytesseract.image_to_string(image, lang="eng")  # Ingliz tilidan
    return text.strip()

def translate_docx_with_images(input_file, target_language="UZ"):
    """Word faylni tarjima qilib, rasm matnlarini ham tarjima qiladi."""
    doc = DocxDocument(input_file)
    translated_doc = DocxDocument()

    for element in doc.element.body:
        if element.tag.endswith('p'):  # Paragraph
            para = doc.paragraphs[doc.element.body.index(element)]
            translated_text = translate_text_with_deepl(para.text, target_language)
            translated_doc.add_paragraph(translated_text)
        elif element.tag.endswith('tbl'):  # Table
            table = doc.tables[doc.element.body.index(element)]
            new_table = translated_doc.add_table(rows=len(table.rows), cols=len(table.columns))
            for i, row in enumerate(table.rows):
                for j, cell in enumerate(row.cells):
                    translated_cell_text = translate_text_with_deepl(cell.text, target_language)
                    new_table.cell(i, j).text = translated_cell_text

    # Rasmlarni qayta ishlash
    for rel in doc.part.rels.values():
        if "image" in rel.target_ref:
            image_data = rel.target_part.blob
            image_stream = BytesIO(image_data)

            # Rasmdagi matnni chiqarish va tarjima qilish
            ocr_text = extract_text_from_image(image_stream)
            translated_text = translate_text_with_deepl(ocr_text, target_language)

            # Tarjima matnini yangi hujjatga rasm bilan qo'shish
            translated_doc.add_paragraph(f"[Image translation]: {translated_text}")
            translated_doc.add_picture(BytesIO(image_data), width=Inches(4))

    output_stream = BytesIO()
    translated_doc.save(output_stream)
    output_stream.seek(0)
    return output_stream

def translate_document_view(request, doc_id):
    doc_obj = get_object_or_404(UploadedFile, id=doc_id)
    translated_file_stream = translate_docx_with_images(doc_obj.file, target_language="UZ")

    response = HttpResponse(
        translated_file_stream,
        content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
    response["Content-Disposition"] = f'attachment; filename="{doc_obj.title}_translated.docx"'
    return response
